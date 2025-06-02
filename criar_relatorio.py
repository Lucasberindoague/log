from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Criar um novo documento
doc = Document()

# Função para adicionar título com formatação
def add_titulo(doc, texto, nivel=1):
    heading = doc.add_heading(texto, level=nivel)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

# Função para adicionar parágrafo com formatação
def add_paragrafo(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# Título do documento
titulo = doc.add_heading('Análise de Chamados de Manutenção - Log', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introdução
add_titulo(doc, '1. Introdução')
add_paragrafo(doc, """
Este relatório apresenta uma análise detalhada dos chamados de manutenção da empresa Log, uma prestadora de serviços especializada em gestão e manutenção de condomínios. O estudo foi realizado com base nos dados extraídos do sistema de gestão de chamados da empresa, consolidados no arquivo "dados-tratados.xlsx".

O objetivo principal desta análise é compreender os padrões, tendências e oportunidades de melhoria no atendimento aos chamados de manutenção, visando otimizar a qualidade do serviço prestado e a satisfação dos clientes. Os dados analisados incluem informações sobre tempo de resolução, avaliação dos clientes, distribuição geográfica, tipos de serviços e outras métricas relevantes.

A análise foi estruturada em nove etapas complementares, cada uma focando em aspectos específicos dos dados, permitindo uma visão abrangente e detalhada do cenário operacional da empresa.
""")

# 2. Análise 1 - Colunas
add_titulo(doc, '2. Análise 1 - Colunas Numéricas e Categóricas')
add_paragrafo(doc, """
Esta etapa inicial focou na análise exploratória das características básicas do conjunto de dados, examinando separadamente as colunas numéricas e categóricas.

Colunas utilizadas:
• Todas as colunas numéricas do dataset
• Todas as colunas categóricas do dataset

Gráficos gerados (pasta graficos_etapa1):
• numerica_[nome_coluna].png: Histogramas e boxplots para cada coluna numérica
• categorica_[nome_coluna].png: Gráficos de barras para cada coluna categórica

Principais observações:
• Identificação da estrutura do dataset com análise de tipos de dados
• Análise de valores ausentes e sua distribuição
• Identificação de outliers nas variáveis numéricas
• Distribuição das categorias nas variáveis categóricas
""")

# 3. Análise 2 - Temporal
add_titulo(doc, '3. Análise 2 - Análise Temporal')
add_paragrafo(doc, """
A segunda etapa focou na análise da distribuição temporal dos chamados, identificando padrões e sazonalidades.

Colunas utilizadas:
• dat_criacao: Data de criação do chamado
• dat_resolucao: Data de resolução do chamado
• des_tipo_servico: Tipo de serviço
• des_status: Status do chamado

Gráficos gerados (pasta graficos_etapa2):
• linha_chamados_por_mes.png
• barras_chamados_por_dia_semana.png
• heatmap_hora_dia.png
• linha_tempo_medio_mensal.png

Principais observações:
• Identificação de padrões sazonais nos chamados
• Análise de picos de demanda por período
• Distribuição dos chamados ao longo da semana
• Evolução temporal do tempo de atendimento
""")

# 4. Análise 3 - Regional
add_titulo(doc, '4. Análise 3 - Análise Regional')
add_paragrafo(doc, """
Esta etapa analisou a distribuição geográfica dos chamados e as características específicas de cada região.

Colunas utilizadas:
• cod_uf: Estado do chamado
• des_condominio: Descrição/localização do condomínio
• dat_criacao: Data de criação
• dat_resolucao: Data de resolução

Gráficos gerados (pasta graficos_etapa3):
• barras_chamados_por_estado.png
• boxplot_tempo_por_estado.png
• heatmap_servico_estado.png
• dispersao_volume_tempo_estado.png

Principais observações:
• Distribuição desigual de chamados entre estados
• Variações regionais no tempo de atendimento
• Relação entre volume e eficiência por estado
• Características específicas de cada região
""")

# 5. Análise 4 - Tempo de Resolução
add_titulo(doc, '5. Análise 4 - Análise do Tempo de Resolução')
add_paragrafo(doc, """
A quarta etapa focou na análise detalhada dos tempos de resolução dos chamados.

Colunas utilizadas:
• dat_criacao: Data de criação
• dat_resolucao: Data de resolução
• des_tipo_servico: Tipo de serviço
• des_prioridade: Prioridade do chamado
• des_status: Status do chamado

Gráficos gerados (pasta graficos_etapa4):
• histograma_tempo_resolucao.png
• boxplot_tempo_servico.png
• linha_tempo_medio_mensal.png
• dispersao_volume_tempo.png

Principais observações:
• Distribuição dos tempos de resolução
• Identificação de outliers
• Relação entre tipo de serviço e tempo
• Evolução temporal da eficiência
""")

# 6. Análise 5 - Avaliação do Cliente
add_titulo(doc, '6. Análise 5 - Análise da Avaliação do Cliente')
add_paragrafo(doc, """
Esta etapa analisou as avaliações feitas pelos clientes após o atendimento.

Colunas utilizadas:
• vlr_nota_avaliacao_cliente_atendimento: Nota do cliente
• cod_uf: Estado
• des_tipo_servico: Tipo de serviço
• tempo_resolucao_dias: Tempo de resolução
• des_prioridade: Prioridade
• des_comentario: Comentário do cliente

Gráficos gerados (pasta graficos_etapa5):
• barras_distribuicao_notas.png
• barras_avaliacao_media_estado.png
• barras_avaliacao_media_servico.png
• dispersao_tempo_avaliacao.png

Principais observações:
• Distribuição das notas de avaliação
• Relação entre tempo e satisfação
• Variações por tipo de serviço
• Análise de comentários dos clientes
""")

# 7. Análise 6 - Categorias
add_titulo(doc, '7. Análise 6 - Análise de Categorias')
add_paragrafo(doc, """
A sexta etapa focou na análise das diferentes categorias de serviços e suas características.

Colunas utilizadas:
• des_tipo_servico: Tipo de serviço
• des_assunto: Assunto do chamado
• des_classificacao: Classificação
• cod_uf: Estado
• vlr_nota_avaliacao_cliente_atendimento: Nota do cliente

Gráficos gerados (pasta graficos_etapa6):
• barras_distribuicao_servicos.png
• heatmap_servico_estado.png
• heatmap_sazonalidade_servico.png
• dispersao_complexidade_servico.png

Principais observações:
• Distribuição dos tipos de serviço
• Padrões de demanda por categoria
• Relação entre categoria e avaliação
• Complexidade por tipo de serviço
""")

# 8. Análise 7 - Não Resolvidos
add_titulo(doc, '8. Análise 7 - Análise de Chamados Não Resolvidos')
add_paragrafo(doc, """
Esta etapa analisou especificamente os chamados que não foram resolvidos.

Colunas utilizadas:
• dat_criacao: Data de criação
• dat_resolucao: Data de resolução
• des_tipo_servico: Tipo de serviço
• des_prioridade: Prioridade
• des_status: Status
• des_assunto: Assunto

Gráficos gerados (pasta graficos_etapa7):
• pizza_status_chamados.png
• barras_nao_resolvidos_servico.png
• linha_tempo_nao_resolvidos.png
• heatmap_nao_resolvidos_estado.png

Principais observações:
• Proporção de chamados não resolvidos
• Características comuns dos não resolvidos
• Distribuição por tipo de serviço
• Tempo em aberto dos chamados
""")

# 9. Análise 8 - Correlações
add_titulo(doc, '9. Análise 8 - Análise de Correlações')
add_paragrafo(doc, """
A oitava etapa focou na análise das correlações entre diferentes variáveis.

Colunas utilizadas:
• tempo_resolucao_dias: Tempo de resolução
• vlr_nota_avaliacao_cliente_atendimento: Nota do cliente
• cod_uf: Estado
• des_tipo_servico: Tipo de serviço

Gráficos gerados (pasta graficos_etapa8):
• matriz_correlacao.png
• scatter_tempo_avaliacao.png
• evolucao_temporal.png
• clusters.png

Principais observações:
• Correlações entre tempo e avaliação
• Padrões temporais
• Clusters de desempenho
• Relações entre variáveis
""")

# 10. Análise 9 - Exploração Livre
add_titulo(doc, '10. Análise 9 - Exploração Livre')
add_paragrafo(doc, """
A última etapa realizou uma análise exploratória livre, buscando insights adicionais.

Colunas utilizadas:
• Todas as colunas relevantes do dataset

Gráficos gerados (pasta graficos_etapa9):
• complexidade_distribuicao.png
• heatmap_hora_dia.png
• matriz_transicao_servicos.png
• clusters_condominios.png
• evolucao_eficiencia.png

Principais observações:
• Análise de complexidade dos chamados
• Padrões cíclicos
• Dependência entre serviços
• Perfis de condomínio
• Eficiência operacional
""")

# 11. Conclusão Geral
add_titulo(doc, '11. Conclusão Geral')
add_paragrafo(doc, """
Principais padrões identificados:
• Concentração significativa de chamados em determinados tipos de serviço
• Variações importantes no tempo de resolução entre diferentes estados
• Correlação fraca entre tempo de resolução e avaliação do cliente
• Padrões sazonais claros em alguns tipos de serviço
• Existência de clusters distintos de desempenho

Hipóteses e interpretações:
• A satisfação do cliente parece estar mais relacionada à qualidade do atendimento do que ao tempo de resolução
• Existe uma oportunidade significativa de padronização de processos entre estados
• A complexidade dos chamados varia significativamente por tipo de serviço
• Alguns condomínios apresentam padrões recorrentes de chamados

Recomendações:
1. Implementar sistema de priorização baseado na complexidade identificada
2. Desenvolver programas de treinamento específicos para estados com menor desempenho
3. Criar protocolos padronizados para os tipos de serviço mais frequentes
4. Estabelecer metas de tempo de resolução específicas por categoria
5. Implementar programa de manutenção preventiva baseado nos padrões identificados
""")

# Salvar o documento
doc.save('relatorio_intermediario.docx')
print("Relatório intermediário gerado com sucesso!")